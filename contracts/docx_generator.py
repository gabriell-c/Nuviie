"""Geração de DOCX fiel ao MODELO CONTRATO NUVIIE.pdf"""

from __future__ import annotations

from django.conf import settings
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Mm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .nuviie_template import NUVIIE_BLUE, render_sections

LOGO_PATH = settings.BASE_DIR / 'static' / 'imgs' / 'logo_branca.png'


def _set_cell_shading(cell, hex_color: str):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), hex_color.lstrip('#'))
    cell._tc.get_or_add_tcPr().append(shading)


def _add_header_table(doc: Document):
    """Logo + barra azul no topo (igual ao PDF)."""
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    logo_cell = table.rows[0].cells[0]
    bar_cell = table.rows[0].cells[1]

    logo_cell.width = Mm(20)
    bar_cell.width = Mm(170)
    _set_cell_shading(logo_cell, NUVIIE_BLUE.lstrip('#'))
    _set_cell_shading(bar_cell, NUVIIE_BLUE.lstrip('#'))

    if LOGO_PATH.exists():
        p = logo_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(LOGO_PATH), width=Mm(14))

    doc.add_paragraph('')


def generate_nuviie_contract_docx(filled_data: dict, output_path: str, contract_title: str):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    _add_header_table(doc)

    for section in render_sections(filled_data):
        text = section.get('text', '').strip()
        if not text:
            continue
        st = section.get('style', 'body')

        if st == 'title':
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(text)
            run.bold = True
            run.font.size = Pt(13)
        elif st in ('party_label', 'clause'):
            para = doc.add_paragraph(text)
            for run in para.runs:
                run.bold = True
        else:
            para = doc.add_paragraph(text)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if st == 'body' else WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph('')
    doc.add_paragraph('_______________________________________')
    doc.add_paragraph(filled_data.get('assinatura_contratante', 'CONTRATANTE'))
    doc.add_paragraph('')
    doc.add_paragraph('_______________________________________')
    doc.add_paragraph('Gabriel Cardoso - Nuviie')

    doc.save(output_path)
